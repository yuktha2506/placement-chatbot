import json
import re
import shutil
import subprocess
import sys


def normalize(text):
    lines = []
    for line in (text or "").replace("\r", "\n").split("\n"):
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        if cleaned:
            lines.append(cleaned)
    return "\n".join(lines).strip()


def extract_pdf(path):
    from pypdf import PdfReader

    parts = []
    reader = PdfReader(path)
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def extract_docx(path):
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_txt(path):
    with open(path, "rb") as handle:
        data = handle.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="ignore")


def attempt_ocr(path, extension):
    result = {
        "text": "",
        "ocrAttempted": False,
        "ocrAvailable": False,
        "ocrError": ""
    }

    tesseract = shutil.which("tesseract")
    if not tesseract:
        result["ocrError"] = "Tesseract OCR binary is not installed."
        return result

    result["ocrAvailable"] = True
    result["ocrAttempted"] = True

    try:
        if extension == ".pdf":
            result["ocrError"] = "OCR for PDF images requires a PDF-to-image renderer, which is not installed."
            return result

        completed = subprocess.run(
            [tesseract, path, "stdout"],
            check=True,
            capture_output=True,
            text=True,
            timeout=90,
        )
        result["text"] = completed.stdout
    except Exception as exc:
        result["ocrError"] = str(exc)

    return result


def main():
    file_path = sys.argv[1]
    extension = sys.argv[2].lower()
    extraction_method = "unknown"
    text = ""
    errors = []

    try:
        if extension == ".pdf":
            extraction_method = "pypdf"
            text = extract_pdf(file_path)
        elif extension == ".docx":
            extraction_method = "python-docx"
            text = extract_docx(file_path)
        elif extension in (".txt", ".md", ".csv"):
            extraction_method = "text"
            text = extract_txt(file_path)
        else:
            extraction_method = "text-fallback"
            text = extract_txt(file_path)
    except Exception as exc:
        errors.append(f"direct extraction failed: {exc}")

    normalized_text = normalize(text)
    ocr = {
        "text": "",
        "ocrAttempted": False,
        "ocrAvailable": False,
        "ocrError": ""
    }

    if len(normalized_text) < 100:
        ocr = attempt_ocr(file_path, extension)
        if ocr["text"]:
            normalized_text = normalize(f"{normalized_text}\n{ocr['text']}")
            extraction_method = f"{extraction_method}+ocr"
        if ocr["ocrError"]:
            errors.append(ocr["ocrError"])

    print(json.dumps({
        "text": normalized_text,
        "extractionMethod": extraction_method,
        "ocrAttempted": ocr["ocrAttempted"],
        "ocrAvailable": ocr["ocrAvailable"],
        "errors": errors
    }))


if __name__ == "__main__":
    main()
